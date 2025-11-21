"""
文件轉換服務 - 將 Word/PDF 規格書轉換為知識庫 Q&A

支援格式:
- .docx (Microsoft Word)
- .pdf (將來擴展)

工作流程:
1. 上傳文件
2. 解析文件內容
3. 使用 AI 提取 Q&A
4. 人工審核編輯
5. 匯入知識庫
"""

import os
import uuid
import json
import asyncio
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime
from docx import Document
import openai
import asyncpg
from asyncpg.pool import Pool


class DocumentConverterService:
    # OpenAI 模型的 context 限制（tokens）
    MODEL_CONTEXT_LIMITS = {
        'gpt-4o': 128000,
        'gpt-4o-mini': 128000,
        'gpt-4-turbo': 128000,
        'gpt-4': 8192,
        'gpt-3.5-turbo': 16385
    }

    def __init__(self, db_pool: Optional[Pool] = None):
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        # 規格書轉換專用模型（需要更強的理解能力和大 context）
        self.model = os.getenv('DOCUMENT_CONVERTER_MODEL', os.getenv('KNOWLEDGE_GEN_MODEL', 'gpt-4o'))
        self.temp_dir = Path('/tmp/document_converter')
        self.temp_dir.mkdir(exist_ok=True)
        self.db_pool = db_pool

        # 轉換任務緩存 (生產環境應使用 Redis)
        self.jobs = {}

        # 意圖快取（減少資料庫查詢）
        self._cached_intents = None

    async def upload_document(self, file_path: str, original_filename: str) -> Dict:
        """
        上傳並驗證文件

        Args:
            file_path: 臨時文件路徑
            original_filename: 原始檔名

        Returns:
            包含 job_id 和文件資訊的字典
        """
        job_id = str(uuid.uuid4())
        file_size = Path(file_path).stat().st_size
        file_ext = Path(original_filename).suffix.lower()

        # 驗證文件格式
        if file_ext not in ['.docx', '.pdf']:
            raise ValueError(f"不支援的檔案格式: {file_ext}。目前只支援 .docx 和 .pdf")

        # 驗證文件大小 (最大 50MB)
        max_size = 50 * 1024 * 1024
        if file_size > max_size:
            raise ValueError(f"檔案過大: {file_size / 1024 / 1024:.1f}MB。最大限制: 50MB")

        # 保存文件
        saved_path = self.temp_dir / f"{job_id}_{original_filename}"
        Path(file_path).rename(saved_path)

        # 創建任務記錄
        self.jobs[job_id] = {
            'job_id': job_id,
            'status': 'uploaded',  # uploaded, parsing, converting, completed, failed
            'file_path': str(saved_path),
            'file_name': original_filename,
            'file_size': file_size,
            'file_type': file_ext[1:],  # 去掉點
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat(),
            'content': None,  # 解析後的純文字
            'qa_list': None,  # AI 提取的 Q&A
            'error': None
        }

        print(f"✅ 文件上傳成功 (job_id: {job_id})")
        print(f"   檔名: {original_filename}")
        print(f"   大小: {file_size / 1024:.1f} KB")
        print(f"   格式: {file_ext}")

        return self.jobs[job_id]

    async def parse_document(self, job_id: str) -> Dict:
        """
        解析文件內容為純文字

        Args:
            job_id: 任務 ID

        Returns:
            包含解析內容的任務資訊
        """
        if job_id not in self.jobs:
            raise ValueError(f"任務不存在: {job_id}")

        job = self.jobs[job_id]
        job['status'] = 'parsing'
        job['updated_at'] = datetime.now().isoformat()

        try:
            file_type = job['file_type']
            file_path = job['file_path']

            if file_type == 'docx':
                content = await self._parse_docx(file_path)
            elif file_type == 'pdf':
                content = await self._parse_pdf(file_path)
            else:
                raise ValueError(f"不支援的檔案格式: {file_type}")

            job['content'] = content
            job['status'] = 'parsed'
            job['updated_at'] = datetime.now().isoformat()

            print(f"✅ 文件解析完成 (job_id: {job_id})")
            print(f"   內容長度: {len(content)} 字元")

            return job

        except Exception as e:
            job['status'] = 'failed'
            job['error'] = str(e)
            job['updated_at'] = datetime.now().isoformat()
            print(f"❌ 文件解析失敗: {e}")
            raise

    async def _parse_docx(self, file_path: str) -> str:
        """
        解析 Word 文件

        Args:
            file_path: 文件路徑

        Returns:
            純文字內容
        """
        doc = Document(file_path)

        # 提取所有段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 只保留非空段落
                paragraphs.append(text)

        # 提取表格內容
        tables_content = []
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    tables_content.append(' | '.join(row_data))

        # 合併段落和表格
        all_content = paragraphs
        if tables_content:
            all_content.append("\n=== 表格內容 ===")
            all_content.extend(tables_content)

        content = '\n'.join(all_content)

        print(f"📄 Word 文件解析:")
        print(f"   段落數: {len(paragraphs)}")
        print(f"   表格數: {len(doc.tables)}")
        print(f"   總字元數: {len(content)}")

        return content

    async def _parse_pdf(self, file_path: str) -> str:
        """
        解析 PDF 文件 (未來實作)

        Args:
            file_path: 文件路徑

        Returns:
            純文字內容
        """
        raise NotImplementedError("PDF 解析功能尚未實作")

    async def convert_to_qa(self, job_id: str, custom_prompt: Optional[str] = None) -> Dict:
        """
        使用 AI 將文件內容轉換為 Q&A

        Args:
            job_id: 任務 ID
            custom_prompt: 自訂提示詞（可選）

        Returns:
            包含 Q&A 列表的任務資訊
        """
        if job_id not in self.jobs:
            raise ValueError(f"任務不存在: {job_id}")

        job = self.jobs[job_id]

        if job['status'] != 'parsed':
            raise ValueError(f"任務狀態錯誤: {job['status']}。請先解析文件")

        job['status'] = 'converting'
        job['updated_at'] = datetime.now().isoformat()

        try:
            content = job['content']

            # 估算 token 並分段處理（考慮 context + TPM 限制）
            # 根據模型動態調整分段大小
            max_context = self.MODEL_CONTEXT_LIMITS.get(self.model, 16385)

            # TPM (Tokens Per Minute) 考量
            # gpt-4o 組織 TPM 限制通常為 30K-90K，保守估計使用 30K
            # 為了避免 rate limit，單次請求應該小於 TPM 限制的 70%
            tpm_limit = 30000 if self.model == 'gpt-4o' else 90000  # gpt-3.5-turbo 通常更高
            safe_request_tokens = int(tpm_limit * 0.7)  # 單次請求安全上限

            # 根據模型容量和 TPM 限制計算安全的分段大小
            # 預留 1000 tokens 給 prompt，4000 tokens 給輸出
            safe_input_tokens = min(max_context - 5000, safe_request_tokens - 4000)
            max_chars = int(safe_input_tokens / 2)  # 中文約 1 字 = 2 tokens

            # 限制範圍：最少 3000 字，最多 10000 字（避免單段太大）
            max_chars = max(3000, min(10000, max_chars))

            print(f"   📏 模型: {self.model} (Context: {max_context}, TPM: ~{tpm_limit})")
            print(f"   📐 分段大小: {max_chars} 字元 (約 {max_chars * 2} tokens)")

            content_chunks = self._split_content(content, max_chars)

            print(f"🤖 開始 AI 轉換 (job_id: {job_id})")
            print(f"   內容分為 {len(content_chunks)} 段處理")
            print(f"   使用模型: {self.model}")

            # 計算 TPM 限制下的安全延遲
            # gpt-4o: 30K TPM，每段約 20K tokens，需要等待 40 秒避免超限
            if len(content_chunks) > 1:
                estimated_tokens_per_chunk = max_chars * 2 + 4000  # 輸入 + 輸出
                delay_seconds = int((estimated_tokens_per_chunk / tpm_limit) * 60 * 1.2)  # 加 20% 緩衝
                delay_seconds = max(20, min(60, delay_seconds))  # 限制在 20-60 秒之間
                print(f"   ⏱️  每段間隔: {delay_seconds} 秒 (避免 TPM 超限)")

            all_qa = []
            for i, chunk in enumerate(content_chunks, 1):
                print(f"   處理第 {i}/{len(content_chunks)} 段...")
                qa_list = await self._call_openai_extract_qa(chunk, custom_prompt)

                # 為每個 Q&A 推薦意圖
                if self.db_pool and qa_list:
                    print(f"   📌 為 {len(qa_list)} 個 Q&A 推薦意圖...")
                    for qa in qa_list:
                        recommended_intent = await self._recommend_intent_for_qa(qa)
                        qa['recommended_intent'] = recommended_intent

                        if recommended_intent['intent_id']:
                            print(f"      ✅ {qa['question_summary'][:30]}... → {recommended_intent['intent_name']} ({recommended_intent['confidence']:.2f})")
                        else:
                            print(f"      ⚠️  {qa['question_summary'][:30]}... → 未分類")

                all_qa.extend(qa_list)

                # 在分段之間添加延遲以避免超過 TPM 限制
                if i < len(content_chunks) and len(content_chunks) > 1:
                    print(f"   ⏳ 等待 {delay_seconds} 秒後處理下一段...")
                    await asyncio.sleep(delay_seconds)

            job['qa_list'] = all_qa
            job['status'] = 'completed'
            job['updated_at'] = datetime.now().isoformat()

            print(f"✅ AI 轉換完成")
            print(f"   提取到 {len(all_qa)} 個 Q&A")
            if self.db_pool:
                intent_recommended = sum(1 for qa in all_qa if qa.get('recommended_intent', {}).get('intent_id'))
                print(f"   已推薦意圖: {intent_recommended}/{len(all_qa)} 個 Q&A")

            return job

        except Exception as e:
            job['status'] = 'failed'
            job['error'] = str(e)
            job['updated_at'] = datetime.now().isoformat()
            print(f"❌ AI 轉換失敗: {e}")
            raise

    def _split_content(self, content: str, max_chars: int) -> List[str]:
        """
        將長文本分段

        Args:
            content: 原始文本
            max_chars: 每段最大字元數

        Returns:
            分段後的文本列表
        """
        if len(content) <= max_chars:
            return [content]

        chunks = []
        paragraphs = content.split('\n')
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para_length = len(para) + 1  # +1 for newline

            if current_length + para_length > max_chars and current_chunk:
                # 當前段落會超過限制，先保存當前chunk
                chunks.append('\n'.join(current_chunk))
                current_chunk = [para]
                current_length = para_length
            else:
                current_chunk.append(para)
                current_length += para_length

        # 保存最後一段
        if current_chunk:
            chunks.append('\n'.join(current_chunk))

        return chunks

    async def _call_openai_extract_qa(self, content: str, custom_prompt: Optional[str] = None) -> List[Dict]:
        """
        呼叫 OpenAI API 提取 Q&A

        Args:
            content: 文件內容
            custom_prompt: 自訂提示詞

        Returns:
            Q&A 列表
        """
        if custom_prompt:
            prompt = custom_prompt.format(content=content)
        else:
            # 詳細 prompt 提供更好的指導（GPT-4o 可以處理）
            prompt = f"""請從以下技術規格書中提取實用的使用者問答對（Q&A）。

**重要：請盡可能提取所有有價值的 Q&A，不要遺漏任何重要功能或操作說明。**

## 提取要求：

1. **問題類型**：
   - 操作步驟類（如何使用某功能？）
   - 規定說明類（什麼情況下需要...？）
   - 功能介紹類（這個功能是做什麼的？）
   - 疑難排解類（遇到XX問題怎麼辦？）

2. **答案品質**：
   - 清晰、具體、完整
   - 包含必要的步驟說明
   - 使用條列式說明（如果適用）
   - 保留重要的細節和注意事項

3. **關鍵字提取**：
   - 從問答中提取 3-5 個關鍵詞
   - 包含專有名詞、功能名稱、操作動作等

## 輸出格式：

請以 JSON 陣列格式輸出，每個 Q&A 包含以下欄位：
- question_summary: 問題摘要（10-30字）
- content: 完整答案（包含所有必要資訊）
- keywords: 關鍵字陣列（3-5個詞）

範例：
[
  {{
    "question_summary": "如何申請停車位？",
    "content": "申請停車位的步驟如下：\\n1. 填寫停車位申請表\\n2. 提供車輛行照影本\\n3. 繳交保證金 5,000 元\\n4. 等待管理中心審核通知",
    "keywords": ["停車位", "申請", "行照", "保證金"]
  }}
]

## 規格書內容：

{content}

請只返回 JSON 格式的輸出，不要包含其他說明文字。"""

        try:
            client = openai.OpenAI(api_key=self.openai_api_key)

            # 計算安全的 max_tokens
            # 估算輸入 tokens（中文約 1 字 = 2 tokens，包含 system + prompt + content）
            estimated_input_tokens = len(content) * 2 + 1000  # +1000 for system and prompt

            # 根據模型動態計算可用的輸出 tokens
            # gpt-4o: 128K context, gpt-4: 8K context, gpt-4-turbo: 128K context
            max_context = self.MODEL_CONTEXT_LIMITS.get(self.model, 16385)  # 預設 16K

            # 計算可用的輸出 tokens（保留 10% 緩衝）
            available_output_tokens = int((max_context - estimated_input_tokens) * 0.9)

            # 限制輸出範圍：最少 1000，最多 4000
            safe_max_tokens = max(1000, min(4000, available_output_tokens))

            print(f"   📊 Token 估算: 輸入 ~{estimated_input_tokens}, 輸出上限 {safe_max_tokens}")

            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "你是一個專業的知識庫管理專家，擅長從技術規格書中提取實用的Q&A。請仔細分析文件內容，提取對使用者有實際幫助的問答對。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=safe_max_tokens  # 設置動態計算的安全上限
            )

            result_text = response.choices[0].message.content.strip()

            # 嘗試解析 JSON
            # 移除可能的 markdown code block 標記
            if result_text.startswith('```'):
                result_text = result_text.split('```')[1]
                if result_text.startswith('json'):
                    result_text = result_text[4:]

            qa_list = json.loads(result_text)

            # 驗證格式
            for qa in qa_list:
                if not all(k in qa for k in ['question_summary', 'content', 'keywords']):
                    raise ValueError(f"Q&A 格式錯誤: {qa}")

                # 確保 keywords 是列表
                if isinstance(qa['keywords'], str):
                    qa['keywords'] = [k.strip() for k in qa['keywords'].split(',')]

            return qa_list

        except json.JSONDecodeError as e:
            print(f"⚠️ JSON 解析失敗: {e}")
            print(f"原始回應: {result_text[:500]}")
            # 返回空列表而不是拋出錯誤
            return []
        except Exception as e:
            print(f"❌ OpenAI API 呼叫失敗: {e}")
            raise

    async def update_qa_list(self, job_id: str, qa_list: List[Dict]) -> Dict:
        """
        更新 Q&A 列表（人工編輯後）

        Args:
            job_id: 任務 ID
            qa_list: 更新後的 Q&A 列表

        Returns:
            更新後的任務資訊
        """
        if job_id not in self.jobs:
            raise ValueError(f"任務不存在: {job_id}")

        job = self.jobs[job_id]
        job['qa_list'] = qa_list
        job['updated_at'] = datetime.now().isoformat()

        print(f"✅ Q&A 列表已更新 (job_id: {job_id})")
        print(f"   Q&A 數量: {len(qa_list)}")

        return job

    async def get_job(self, job_id: str) -> Optional[Dict]:
        """
        獲取任務資訊

        Args:
            job_id: 任務 ID

        Returns:
            任務資訊，不存在則返回 None
        """
        return self.jobs.get(job_id)

    async def estimate_cost(self, content_length: int) -> Dict:
        """
        估算轉換成本

        Args:
            content_length: 內容長度（字元數）

        Returns:
            成本估算資訊
        """
        # GPT-4 pricing (approximate)
        # Input: $0.03 / 1K tokens
        # Output: $0.06 / 1K tokens

        # 估算 tokens (中文約 1 字 = 1.5 tokens, 英文約 1 字 = 0.25 tokens)
        # 保守估計使用 1.5
        estimated_tokens = int(content_length * 1.5)

        # 假設輸出是輸入的 50%
        output_tokens = int(estimated_tokens * 0.5)

        # 計算成本
        input_cost = (estimated_tokens / 1000) * 0.03
        output_cost = (output_tokens / 1000) * 0.06
        total_cost = input_cost + output_cost

        return {
            'content_length': content_length,
            'estimated_input_tokens': estimated_tokens,
            'estimated_output_tokens': output_tokens,
            'estimated_cost_usd': round(total_cost, 2),
            'model': self.model
        }

    async def cleanup_job(self, job_id: str):
        """
        清理任務文件

        Args:
            job_id: 任務 ID
        """
        if job_id in self.jobs:
            job = self.jobs[job_id]
            file_path = Path(job['file_path'])

            if file_path.exists():
                file_path.unlink()
                print(f"🗑️  已刪除文件: {file_path}")

            del self.jobs[job_id]
            print(f"✅ 任務已清理 (job_id: {job_id})")

    async def _get_all_intents(self) -> List[Dict]:
        """
        取得所有可用的意圖

        Returns:
            意圖列表，包含 id, name, description
        """
        if not self.db_pool:
            print("   ⚠️  未設定資料庫連接池，跳過意圖載入")
            return []

        try:
            # 使用快取避免重複查詢
            if self._cached_intents is not None:
                return self._cached_intents

            async with self.db_pool.acquire() as conn:
                rows = await conn.fetch("""
                    SELECT id, name, description
                    FROM intents
                    ORDER BY id
                """)

                self._cached_intents = [dict(row) for row in rows]
                print(f"   ✅ 載入 {len(self._cached_intents)} 個意圖")
                return self._cached_intents

        except Exception as e:
            print(f"   ⚠️  載入意圖失敗: {e}")
            return []

    async def _recommend_intent_for_qa(self, qa: Dict) -> Dict:
        """
        為單個 Q&A 推薦意圖

        複製自 knowledge_import_service._recommend_intents() 的邏輯

        Args:
            qa: Q&A 資料，包含 question_summary, content, keywords

        Returns:
            推薦結果，包含 intent_id, intent_name, confidence, reasoning
        """
        try:
            # 1. 取得所有意圖（使用快取）
            intents = await self._get_all_intents()

            if not intents:
                return {
                    'intent_id': None,
                    'intent_name': '未分類',
                    'confidence': 0.0,
                    'reasoning': '系統中沒有可用意圖'
                }

            # 2. 建立意圖清單文字
            intent_list = "\n".join([
                f"- {i['id']}: {i['name']} ({i['description']})"
                for i in intents
            ])

            # 3. 呼叫 LLM 推薦
            keywords_str = ', '.join(qa.get('keywords', []))
            prompt = f"""請根據以下問答內容，從意圖清單中選擇最合適的意圖。

問題：{qa['question_summary']}
答案：{qa['content'][:200]}
關鍵字：{keywords_str}

可用的意圖清單：
{intent_list}

請以 JSON 格式回應：
{{
  "intent_id": 推薦的意圖 ID（數字）,
  "intent_name": 意圖名稱,
  "confidence": 信心度（0.0-1.0）,
  "reasoning": 推薦理由（簡短說明）
}}

只輸出 JSON，不要加其他說明。"""

            client = openai.OpenAI(api_key=self.openai_api_key)
            response = client.chat.completions.create(
                model=self.model,
                temperature=0.3,
                max_tokens=500,  # 意圖推薦只需要小量輸出
                response_format={"type": "json_object"},
                messages=[{"role": "user", "content": prompt}]
            )

            result = json.loads(response.choices[0].message.content)

            return {
                'intent_id': result.get('intent_id'),
                'intent_name': result.get('intent_name'),
                'confidence': result.get('confidence', 0.8),
                'reasoning': result.get('reasoning', '')
            }

        except Exception as e:
            print(f"   ⚠️  意圖推薦失敗: {e}")
            return {
                'intent_id': None,
                'intent_name': '未分類',
                'confidence': 0.0,
                'reasoning': f'推薦失敗: {str(e)}'
            }
